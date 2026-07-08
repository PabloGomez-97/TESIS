"""Genera outline del potencial paper en formato .docx."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "outline_paper_dryspells_chile.docx"


def add_para(doc, text, style=None, bold=False):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p


def add_bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


def build():
    doc = Document()

    # Estilos base
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Portada
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Outline — Potencial Paper")
    run.bold = True
    run.font.size = Pt(18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Changes in Dry-Spell Duration Distributions over Chile:\n"
        "Historical Comparison with CR2MET and Future Projections from ALADIN CHP12"
    )
    run.font.size = Pt(12)
    run.italic = True

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Pablo Gómez — Tesis / Storyboard para publicación\n").bold = True
    meta.add_run("Basado en reunión 30-jun-2026 con Cristian Martinez-Villalobos\n")
    meta.add_run("Referencia estructural: JCLI-D-21-0617 (Martinez-Villalobos, Neelin & Pendergrass 2022)")

    doc.add_page_break()

    # Resumen ejecutivo
    doc.add_heading("Resumen ejecutivo", level=1)
    add_para(
        doc,
        "Historia en una línea: Las distribuciones de duración de dry spells capturan intermitencia "
        "que las métricas agregadas no ven; en Chile, ALADIN (calibrado) reproduce en parte el "
        "histórico respecto a CR2MET, pero proyecta mayor probabilidad de spells largos bajo SSP5-8.5.",
    )
    add_para(
        doc,
        "Convención del documento: cada viñeta numerada [P#] corresponde a UNA oración = UN párrafo "
        "del paper. Las figuras se asignan después de definir el mensaje del párrafo (historia primero, "
        "figura después).",
    )
    add_para(
        doc,
        "Notación: R₀ (antes τ* en la tesis) denota el umbral calibrado en ALADIN (= 5.285 mm global) "
        "para igualar la fracción integrada de wet days de CR2MET @ 1 mm.",
    )

    # ABSTRACT
    doc.add_heading("Abstract", level=1)
    abstract = [
        "[P1] La precipitación media no describe cómo se reparte la lluvia en el tiempo, y periodos secos consecutivos (dry spells) ofrecen información sobre intermitencia y riesgo de sequía relevante para gestión hídrica en Chile.",
        "[P2] Usamos CR2MET v2.5 como referencia histórica y ALADIN CHP12 para el periodo histórico (1980–2014) y futuro SSP5-8.5 (2040–2074), definiendo dry spells tras calibrar el umbral de día húmedo en ALADIN (R₀ = 5.285 mm) para igualar la fracción integrada de wet days de CR2MET a 1 mm.",
        "[P3] Caracterizamos la distribución de duraciones mediante PDFs empíricas (normalizadas para integrar al número de spells) y cuantificamos cambios futuros con risk ratios (RR) y bootstrap por años.",
        "[P4] Históricamente, ALADIN subestima la persistencia media y extrema (p99) de dry spells respecto a CR2MET incluso calibrado, con acuerdo regional variable en las PDFs por cuenca.",
        "[P5] En proyecciones, varias regiones muestran RR > 1 y significativo para spells ≥ 20 días, con señal que se refuerza en la cola de duraciones muy largas; las PDFs por cuenca confirman extensión de la cola derecha en el futuro.",
        "[P6] Concluimos que el análisis distributivo de dry spells complementa métricas tradicionales y sugiere mayor riesgo de persistencia seca futura, con implicancias para evaluación de modelos regionales y planificación adaptativa.",
    ]
    for p in abstract:
        add_bullet(doc, p)

    # SECTION 1
    doc.add_heading("1. Introduction", level=1)

    doc.add_heading("1.1 Motivation: why dry spells matter", level=2)
    s11 = [
        "[P7] Chile ha experimentado megasequías y variabilidad hidrológica donde el total anual de precipitación puede enmascarar cambios en la frecuencia y persistencia de periodos secos.",
        "[P8] Dos climas con la misma precipitación anual pueden diferir radicalmente en intermitencia —p. ej., 300 mm repartidos en 2 eventos vs. 15— con consecuencias distintas para reservorios, agricultura y ecosistemas.",
        "[P9] Las métricas agregadas (media, percentiles de precipitación diaria, índices de sequía basados en acumulados) no caracterizan explícitamente la duración de rachas secas consecutivas.",
        "[P10] Las distribuciones de probabilidad de duración de dry spells ofrecen un marco parsimonioso para describir intermitencia y colas de eventos extremadamente largos, en línea con enfoques recientes para precipitación diaria (Martinez-Villalobos & Neelin 2018; Martinez-Villalobos et al. 2022).",
    ]
    for p in s11:
        add_bullet(doc, p)
    add_para(doc, "→ Fig. 1 (CREAR): esquema conceptual — precipitación total similar vs. distinta intermitencia / persistencia de dry spells (motivación, no datos).", bold=True)

    doc.add_heading("1.2 Regional climate models and the comparison problem", level=2)
    s12 = [
        "[P11] Los productos gridded de referencia (p. ej. CR2MET) y los modelos climáticos regionales (RCM) como ALADIN CHP12 no son directamente intercambiables sin considerar cómo cada producto define y simula días húmedos.",
        "[P12] Comparar CR2MET y ALADIN con el mismo umbral fijo (1 mm/día) introduce sesgos sistemáticos en frecuencia e intensidad de wet days que contaminarían cualquier análisis de dry spells.",
        "[P13] Por ello, el sesgo de umbral se trata aquí como paso metodológico (calibración R₀), no como la pregunta científica central del paper.",
    ]
    for p in s12:
        add_bullet(doc, p)
    add_para(doc, "→ Fig. S1 (EXISTE, suplemento): fig_p3_delta_wetdays.png — delta fracción wet days ALADIN−CR2MET @ 1 mm.", bold=True)

    doc.add_heading("1.3 Objectives and paper structure", level=2)
    s13 = [
        "[P14] Objetivo 1: evaluar históricamente (1980–2014) cómo ALADIN representa la distribución de duraciones de dry spells respecto a CR2MET en Chile continental.",
        "[P15] Objetivo 2: cuantificar cambios proyectados (SSP5-8.5, 2040–2074) en la probabilidad de dry spells largos mediante risk ratios y PDFs regionales por cuenca.",
        "[P16] Objetivo 3: documentar la sensibilidad de las PDFs al criterio de umbral y justificar el uso de R₀ calibrado para comparaciones equitativas.",
        "[P17] El resto del paper se organiza: datos y métodos (§2), evaluación histórica (§3), proyecciones futuras (§4), y discusión (§5).",
    ]
    for p in s13:
        add_bullet(doc, p)
    add_para(doc, "→ Fig. 2 (CREAR o adaptar): mapa del dominio — Chile continental, cuencas (Loa, Maipo, Maule, Biobío) y regiones administrativas usadas en RR.", bold=True)

    # SECTION 2
    doc.add_heading("2. Data and Methods", level=1)

    doc.add_heading("2.1 Datasets, domain, and periods", level=2)
    s21 = [
        "[P18] Usamos CR2MET v2.5 (mm/día, grilla lat/lon) como referencia observacional para el histórico 1980–2014.",
        "[P19] Usamos ALADIN CHP12 en su grilla nativa (~3587 celdas sobre Chile), con precipitación convertida de kg m⁻² s⁻¹ a mm/día.",
        "[P20] CR2MET se interpola linealmente a los centros de celda de ALADIN para comparaciones pixel a pixel.",
        "[P21] Para el futuro empleamos ALADIN bajo escenario SSP5-8.5 para 2040–2074 (35 años), emparejado con el histórico ALADIN 1980–2014 del mismo largo.",
        "[P22] El dominio principal es Chile continental; los análisis regionales usan cuencas hidrográficas (Loa, Maipo, Maule, Biobío) y regiones administrativas para los risk ratios.",
    ]
    for p in s21:
        add_bullet(doc, p)
    add_para(doc, "→ Tabla 1 (CREAR): resumen de datasets, resolución, periodos y variables.", bold=True)

    doc.add_heading("2.2 Dry spells and wet-day definition", level=2)
    s22 = [
        "[P23] Definimos wet day como día con precipitación diaria ≥ umbral R (mm/día); dry day como R estrictamente menor.",
        "[P24] Un dry spell es una racha consecutiva de dry days; su duración es el número de días de la racha.",
        "[P25] Para comparaciones históricas equitativas usamos el criterio i (global): CR2MET con R = 1 mm y ALADIN con R₀ = 5.285 mm, umbral que iguala la fracción espacial media de wet days de CR2MET @ 1 mm (Martinez-Villalobos et al. 2022, JCLI-D-21-0590).",
        "[P26] Denotamos R₀ (antes τ* en la tesis) para no confundir con el exponente τ_P de leyes de potencia en PDFs de precipitación (JCLI-D-21-0617).",
        "[P27] Análisis de sensibilidad usan criterio ii (mismo umbral): 1 mm en ambos; y criterio iii (local): R₀(x,y) por píxel.",
    ]
    for p in s22:
        add_bullet(doc, p)
    add_para(doc, "→ Fig. 3 (EXISTE): fig_p6_ftau.png — curva F(R) de ALADIN; punto CR2MET @ 1 mm; cuadrado ALADIN @ R₀ = 5.285 mm.", bold=True)

    doc.add_heading("2.3 Calibrated threshold R₀", level=2)
    s23 = [
        "[P28] La calibración busca R₀ tal que el promedio espacial sobre Chile de la fracción temporal de días con pr ≥ R en ALADIN iguale la de CR2MET con 1 mm.",
        "[P29] CR2MET @ 1 mm da F = 20.94%; ALADIN @ 1 mm da F = 30.68% (+9.7 pp); ALADIN @ R₀ = 5.285 mm reproduce F = 20.94% con error < 0.0001 pp.",
        "[P30] La variante espacial R₀(x,y) permite calibración local por píxel, reduciendo residuos espaciales pero con mayor complejidad interpretativa.",
    ]
    for p in s23:
        add_bullet(doc, p)
    add_para(doc, "→ Fig. S2 (EXISTE, suplemento): tau_local_map.png / tau_integrado_map.png — mapa R₀ local vs. integrado uniforme.", bold=True)

    doc.add_heading("2.4 PDF construction and normalization", level=2)
    s24 = [
        "[P31] Construimos PDFs empíricas de duración t con bins aproximadamente uniformes en espacio log(t), siguiendo la convención de PDFs de precipitación en escala log-log.",
        "[P32] Normalizamos para que la integral de la PDF respecto a t iguale el número total de dry spells (n_spells), no 1; así se preserva si el futuro tiene más o menos eventos además de cambios de forma.",
        "[P33] Esto distingue la PDF (distribución de duraciones de eventos) de una tendencia temporal (evolución año a año), conceptos que no deben confundirse.",
        "[P34] Reportamos media, mediana, p90, p99 y máximo del pool regional de spells por cuenca/periodo como complemento escalar a las PDFs.",
    ]
    for p in s24:
        add_bullet(doc, p)
    add_para(doc, "→ Ecuación (1): definición de PDF con ∫₀^∞ PDF(t) dt = n_spells.", bold=True)

    doc.add_heading("2.5 Regional pooling and spatial statistics", level=2)
    s25 = [
        "[P35] Para PDFs por cuenca, agrupamos (pool) todos los dry spells de todos los píxeles de la cuenca en un único conjunto de eventos, siguiendo el enfoque regional de Martinez-Villalobos & Neelin (2018).",
        "[P36] Para mapas espaciales, calculamos por píxel la duración media de spells y el percentil 99 (t99) de duraciones, y reportamos diferencias ALADIN − CR2MET.",
        "[P37] La significancia espacial de diferencias históricas se evalúa con bootstrap por años (re-muestreo de años, no de días).",
    ]
    for p in s25:
        add_bullet(doc, p)

    doc.add_heading("2.6 Risk ratios and future changes", level=2)
    s26 = [
        "[P38] Definimos el risk ratio RR(D) = P_fut(t ≥ D) / P_hist(t ≥ D), probabilidad condicional de spells de al menos D días en futuro vs. histórico.",
        "[P39] Consideramos spells con inicio entre marzo y noviembre para enfocar la estación seca/húmeda relevante.",
        "[P40] Estimamos intervalos de confianza al 95% mediante bootstrap por años sobre el pool regional de cada región administrativa.",
        "[P41] RR > 1 con IC95 completamente por encima de 1 indica aumento significativo de la probabilidad de spells largos.",
    ]
    for p in s26:
        add_bullet(doc, p)
    add_para(doc, "→ Ecuación (2): definición de RR(D).", bold=True)

    doc.add_heading("2.7 Sensitivity to threshold choice", level=2)
    s27 = [
        "[P42] Para demostrar que las conclusiones dependen del umbral, comparamos cuatro series en cuencas seleccionadas: CR2MET @ 1 mm, ALADIN @ 1 mm, ALADIN @ R₀ global, ALADIN @ R₀ local.",
        "[P43] Sin calibración (ALADIN @ 1 mm), ALADIN produce muchos más spells más cortos —p. ej. Loa: media 27 d vs. 66 d en CR2MET— distorsionando la comparación.",
        "[P44] Con R₀ global o local, las PDFs de ALADIN se acercan sustancialmente a CR2MET, justificando el criterio i para el análisis principal.",
    ]
    for p in s27:
        add_bullet(doc, p)
    add_para(doc, "→ Fig. 4 (EXISTE): pdf_calibracion_Loa_logy.png — panel 4 curvas (Loa).", bold=True)
    add_para(doc, "→ Fig. S3 (EXISTE, suplemento): calibración Maule y Biobío.", bold=True)

    # SECTION 3
    doc.add_heading("3. Historical evaluation (1980–2014)", level=1)

    doc.add_heading("3.1 CR2MET climatology of dry-spell persistence", level=2)
    s31 = [
        "[P45] CR2MET muestra fuerte contraste norte–sur: en el norte árido (cuenca Loa) la duración media de dry spells supera 50 días, mientras en el sur húmedo (Biobío) la media es ~6 días.",
        "[P46] Los percentiles altos (p99) alcanzan cientos de días en zonas áridas, reflejando spells extremadamente largos en el desierto.",
        "[P47] CR2MET también muestra sequedización moderada reciente (1980–2000 vs. 2001–2021) en fracción de wet days y precipitación media, contexto independiente del sesgo del RCM.",
    ]
    for p in s31:
        add_bullet(doc, p)
    add_para(doc, "→ Fig. S4 (EXISTE, suplemento): fig_delta_cr2met_wetdays.png — delta wet days CR2MET P2−P1.", bold=True)

    doc.add_heading("3.2 ALADIN vs. CR2MET: spatial persistence biases", level=2)
    s32 = [
        "[P48] Con criterio i (R₀ calibrado), ALADIN subestima la duración media espacial de dry spells: CR2MET 42.5 d vs. ALADIN 35.1 d (Δ = −7.9 d sobre Chile).",
        "[P49] La subestimación es aún mayor sin calibración (criterio ii, 1 mm ambos): Δ = −27.5 d, mostrando que parte del sesgo de persistencia se debe al umbral y parte a la simulación del modelo.",
        "[P50] Para spells extremadamente largos (t99), CR2MET da 221.8 d vs. ALADIN 178.2 d (Δ = −44.7 d); ~65% de celdas con ALADIN significativamente más corto.",
        "[P51] La calibración local (criterio iii) reduce el sesgo de t99 (Δ = −35.3 d) pero no lo elimina, indicando limitaciones estructurales del RCM en persistencia seca.",
    ]
    for p in s32:
        add_bullet(doc, p)
    add_para(doc, "→ Fig. 5 (GENERAR): mapas espaciales duración media y/o t99 — exportar desde pregunta7_mapas_dryspells.ipynb.", bold=True)

    doc.add_heading("3.3 Basin-scale PDFs: agreement and regional differences", level=2)
    s33 = [
        "[P52] Las PDFs por cuenca (pool regional, criterio i, normalización a n_spells) permiten comparar forma y cola de la distribución entre CR2MET hist, ALADIN hist y la referencia futura.",
        "[P53] En Loa (árido), CR2MET tiene media 65.6 d; ALADIN hist 54.6 d; las tres curvas en log-log muestran ALADIN con menos probabilidad en la cola de spells muy largos.",
        "[P54] En Maipo y Maule (centro), el acuerdo mejora: medias ~8–9 d (CR2MET) vs. ~7–8 d (ALADIN hist), con PDFs casi paralelas en escala log-log.",
        "[P55] En Biobío (sur), medias ~6.2 d (CR2MET) vs. ~5.3 d (ALADIN hist); las distribuciones son más estrechas y la diferencia relativa menor que en el norte.",
        "[P56] En conjunto, ALADIN reproduce razonablemente la forma de las PDFs en cuencas centrales y meridionales, pero subestima persistencia en el norte y en la cola extrema.",
    ]
    for p in s33:
        add_bullet(doc, p)
    add_para(doc, "→ Fig. 6 (EXISTE): pdf_Loa_i_global.png — PDF Loa.", bold=True)
    add_para(doc, "→ Fig. 7 (EXISTE): pdf_Maule_i_global.png — PDF Maule.", bold=True)
    add_para(doc, "→ Tabla 2 (EXISTE): fig_p9_tabla_resumen.png — estadísticas por cuenca.", bold=True)

    doc.add_heading("3.4 Historical trends in ALADIN (null result)", level=2)
    s34 = [
        "[P57] Analizamos tendencias lineales anuales en duración media y t99 de dry spells en ALADIN histórico 1980–2014 con R₀ = 5.285 mm.",
        "[P58] La pendiente media espacial es ~−0.027 d/década para duración media y ~+0.001 d/década para t99 — prácticamente cero.",
        "[P59] Menos del 0.2% de las celdas muestra tendencia significativa (p < 0.05), indicando ausencia de tendencia robusta en persistencia de dry spells en ALADIN histórico.",
        "[P60] Esto contrasta con la sequedización en CR2MET (P1 vs. P2) y sugiere que la señal de cambio en persistencia seca no es detectable en ALADIN histórico con este diseño, reforzando el foco en proyecciones futuras para el mensaje principal.",
    ]
    for p in s34:
        add_bullet(doc, p)
    add_para(doc, "→ Fig. S5 (GENERAR o suplemento): mapa de tendencias ALADIN hist — pregunta 8, opcional.", bold=True)

    # SECTION 4
    doc.add_heading("4. Future projections (SSP5-8.5, 2040–2074)", level=1)

    doc.add_heading("4.1 Risk ratios: more long dry spells", level=2)
    s41 = [
        "[P61] Evaluamos RR(D) para spells con inicio marzo–noviembre, pool por región administrativa, ALADIN hist 1980–2014 vs. futuro 2040–2074.",
        "[P62] A D = 20 días, Coquimbo muestra RR = 1.27 (15.4% → 19.6%), IC95 [1.09, 1.48], aumento significativo.",
        "[P63] O'Higgins tiene RR = 1.30 (5.8% → 7.5%), IC95 [1.04, 1.63], también significativo.",
        "[P64] La Araucanía presenta el mayor RR regional a 20 d: 1.84 (0.74% → 1.37%), IC95 [1.16, 3.12], con base histórica baja pero cambio relativo fuerte.",
        "[P65] Los Lagos muestra RR = 2.24 a 20 d pero no significativo (IC95 cruza 1), dado el muy bajo número histórico de spells largos.",
        "[P66] Las curvas RR vs. D muestran que la señal se refuerza para umbrales mayores (cola de spells muy largos), especialmente en La Araucanía (RR > 3 para D ≥ 30 d).",
    ]
    for p in s41:
        add_bullet(doc, p)
    add_para(doc, "→ Fig. 8 (EXISTE): rr_curves_aladin.png — RR vs. umbral D (4 regiones).", bold=True)

    doc.add_heading("4.2 Basin PDFs: future extension of the tail", level=2)
    s42 = [
        "[P67] Las PDFs por cuenca con normalización a n_spells permiten ver simultáneamente cambios en número de eventos y en forma de la distribución.",
        "[P68] En Loa, ALADIN futuro aumenta la media (52.4 d vs. 49.5 d hist) y el máximo (1853 d vs. 1433 d), con la curva futura extendiéndose hacia spells más largos en la cola.",
        "[P69] En Maipo y Maule, el futuro muestra incrementos en p99 y máximo (p. ej. Maule p99: 65 d hist → 78 d futuro) con PDFs que se desplazan levemente hacia duraciones mayores.",
        "[P70] En Biobío, la media futura (5.7 d) iguala a CR2MET histórico, pero p99 y máximo aumentan (42 d y 153 d futuro vs. 38 d y 137 d hist), consistente con mayor probabilidad de spells largos pese a climatología húmeda.",
        "[P71] En escala log-log, las curvas futuras no son simples traslaciones paralelas: el cambio es más pronunciado en la cola derecha, coherente con los RR para D grandes.",
    ]
    for p in s42:
        add_bullet(doc, p)
    add_para(doc, "→ Fig. 9 (EXISTE): PDF Loa con énfasis en curva ALADIN futuro.", bold=True)
    add_para(doc, "→ Fig. 10 (EXISTE): pdf_Biobio_i_global.png — contraste sur húmedo.", bold=True)

    doc.add_heading("4.3 Consistency between RR and PDF diagnostics", level=2)
    s43 = [
        "[P72] Los RR regionales y las PDFs por cuenca son diagnósticos complementarios: RR resume cambio en probabilidad de excedencia; PDF muestra la forma completa de la distribución.",
        "[P73] Regiones con RR significativo a 20 d (Coquimbo, O'Higgins, La Araucanía) son consistentes con cuencas donde las PDFs muestran extensión de la cola (Maipo/Maule en zona central).",
        "[P74] La no-significancia en Los Lagos refleja alta incertidumbre en eventos raros, ilustrando el límite del bootstrap cuando la probabilidad histórica es < 0.2%.",
    ]
    for p in s43:
        add_bullet(doc, p)

    # SECTION 5
    doc.add_heading("5. Summary and Discussion", level=1)

    doc.add_heading("5.1 Summary of findings", level=2)
    s51 = [
        "[P75] Las PDFs de duración de dry spells proveen información sobre intermitencia y persistencia seca que métricas agregadas de precipitación no capturan.",
        "[P76] La calibración R₀ es necesaria para comparaciones equitativas CR2MET–ALADIN, pero no suficiente: ALADIN aún subestima persistencia histórica, especialmente en la cola.",
        "[P77] ALADIN histórico 1980–2014 no muestra tendencia robusta en persistencia de dry spells.",
        "[P78] Las proyecciones SSP5-8.5 indican aumento significativo de la probabilidad de dry spells largos en varias regiones, con señal creciente para duraciones extremas.",
    ]
    for p in s51:
        add_bullet(doc, p)

    doc.add_heading("5.2 Implications for drought risk and model evaluation", level=2)
    s52 = [
        "[P79] Para gestión hídrica y agricultura, el aumento de probabilidad de spells ≥ 20–30 días puede ser más relevante que cambios modestos en precipitación media.",
        "[P80] La metodología distributiva propuesta es transferible a otros RCMs y dominios, en línea con métricas de PDF de precipitación diaria (JCLI-D-21-0617).",
        "[P81] Evaluar RCMs solo por sesgo en wet days o media de precipitación es insuficiente si el interés está en persistencia seca y eventos extremos de duración.",
    ]
    for p in s52:
        add_bullet(doc, p)

    doc.add_heading("5.3 Limitations and future work", level=2)
    s53 = [
        "[P82] CR2MET no está disponible para el futuro, por lo que el cambio futuro se evalúa solo con ALADIN (hist vs. fut del mismo modelo).",
        "[P83] El pool regional agrega spells de distintos píxeles, mezclando heterogeneidad espacial; análisis por píxel o por subcuenca podría refinar la interpretación.",
        "[P84] La calibración R₀ iguala fracción de wet days, no garantiza igualdad en intensidad, variabilidad interanual ni física de spells.",
        "[P85] Trabajo futuro: extender a más cuencas y escenarios, vincular PDFs de dry spells con índices de sequía operativos, y comparar con otros RCMs de CHP.",
    ]
    for p in s53:
        add_bullet(doc, p)

    doc.add_page_break()

    # Tabla figuras
    doc.add_heading("Inventario de figuras", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "#"
    hdr[1].text = "Mensaje que ilustra"
    hdr[2].text = "Archivo / estado"

    figures = [
        ("1", "Motivación: intermitencia importa", "CREAR (esquema)"),
        ("2", "Dominio y cuencas/regiones", "CREAR"),
        ("3", "Calibración R₀ = 5.285 mm", "fig_p6_ftau.png ✓"),
        ("4", "Sensibilidad al umbral (4 curvas)", "pdf_calibracion_Loa_logy.png ✓"),
        ("5", "Sesgo espacial persistencia hist", "GENERAR desde pregunta 7"),
        ("6", "PDF cuenca árida (Loa)", "pdf_Loa_i_global.png ✓"),
        ("7", "PDF cuenca central (Maule)", "pdf_Maule_i_global.png ✓"),
        ("8", "Risk ratios futuro", "rr_curves_aladin.png ✓"),
        ("9", "Cola futura en árido", "mismo panel Loa, énfasis futuro ✓"),
        ("10", "Cola futura en sur húmedo", "pdf_Biobio_i_global.png ✓"),
        ("Tabla 2", "Estadísticas por cuenca", "fig_p9_tabla_resumen.png ✓"),
        ("S1", "Sesgo wet days @ 1 mm", "fig_p3_delta_wetdays.png ✓"),
        ("S2", "R₀ local vs global", "tau_local_map.png ✓"),
        ("S3", "Calibración Maule/Biobío", "pdf_calibracion_*.png ✓"),
        ("S4", "Sequedización CR2MET P1/P2", "fig_delta_cr2met_wetdays.png ✓"),
        ("S5", "Tendencias ALADIN hist (null)", "pregunta 8, opcional"),
    ]
    for num, msg, status in figures:
        row = table.add_row().cells
        row[0].text = num
        row[1].text = msg
        row[2].text = status

    doc.add_paragraph()
    doc.add_heading("Qué falta antes de escribir el paper", level=1)
    pending = [
        "Exportar mapas de pregunta 7 (duración media, t99, Δ) — Fig. 5.",
        "Mapa de dominio con cuencas y regiones — Fig. 2.",
        "Esquema conceptual de intermitencia — Fig. 1 (simple, 1 panel).",
        "Panel multi-cuenca de calibración (Loa + Maule + Biobío en una figura) — versión publicación de Fig. 4.",
        "Decidir idioma final y nombre definitivo de R₀ con Cristian.",
    ]
    for item in pending:
        add_bullet(doc, item)

    doc.add_heading("Cómo usar esto en la tesis", level=1)
    thesis_table = doc.add_table(rows=1, cols=2)
    thesis_table.style = "Table Grid"
    h = thesis_table.rows[0].cells
    h[0].text = "Paper (conciso)"
    h[1].text = "Tesis (expandida)"
    thesis_rows = [
        ("§1 motivación breve", "Capítulo background + megasequía Chile"),
        ("§2.2–2.3 R₀ en 2 párrafos", "Punto 4 completo + puntos 1–2 en anexo"),
        ("§3.4 tendencias null", "Punto 8 completo"),
        ("Suplemento sesgo", "Capítulo metodológico extendido"),
    ]
    for a, b in thesis_rows:
        r = thesis_table.add_row().cells
        r[0].text = a
        r[1].text = b

    doc.add_paragraph()
    doc.add_heading("Contexto de la reunión (30-jun-2026)", level=1)
    meeting = [
        "Prioridad 2 acordada: generar storyboard/outline de potencial paper (no solo tesis).",
        "Historia principal: cambios en distribución de dry spells (histórico vs. futuro), no sesgo como eje central.",
        "τ* / R₀ es punto metodológico breve; precedente en Martinez-Villalobos et al. 2022 (JCLI-D-21-0590).",
        "PDFs deben integrar a n_spells (no a 1) para preservar diferencias hist vs futuro.",
        "Figuras de calibración: 4 curvas (CR2MET 1 mm, ALADIN 1 mm, ALADIN R₀ global, ALADIN R₀ local) en 2–3 cuencas.",
        "Estructura modelo: JCLI-D-21-0617 — secciones, subsecciones, párrafos encadenados, figuras después del mensaje.",
        "Reunión siguiente mencionada: 10 de julio.",
    ]
    for item in meeting:
        add_bullet(doc, item)

    doc.save(OUT)
    print(f"Guardado: {OUT}")


if __name__ == "__main__":
    build()
